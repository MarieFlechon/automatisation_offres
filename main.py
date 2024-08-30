import json
from docx import Document

# Fonction pour charger les formulaires depuis un fichier JSON centralisé
def charger_formulaires(fichier_json):
    with open(fichier_json, 'r', encoding='utf-8') as f:
        data = json.load(f)  # Charger le JSON qui est une liste de formulaires
    return data  # Retourner directement la liste

# Fonction pour générer un document d'offre à partir des données du formulaire
def generer_offre(formulaire, output_path):
    document = Document()
    
    # Ajouter les sections du document avec les informations du formulaire
    document.add_heading(f'Offre pour {formulaire["client"]}', 0)
    document.add_paragraph(f"Secteur : {formulaire['secteur']}")
    document.add_paragraph(f"Description du projet : {formulaire['description']}")
    document.add_paragraph(f"Budget : {formulaire['budget']}")
    document.add_paragraph(f"Objectifs : {formulaire['objectifs']}")
    document.add_paragraph(f"Délais : {formulaire['délais']}")

    # Enregistrer le document dans le fichier spécifié
    document.save(output_path)

# Charger les formulaires depuis le fichier JSON
formulaires = charger_formulaires('formulaires.json')

# Parcourir les formulaires et générer une offre pour chaque client
for formulaire in formulaires:
    output_file = f"Offre_{formulaire['client']}.docx"
    generer_offre(formulaire, output_file)
    print(f"Offre générée pour {formulaire['client']} dans le fichier {output_file}")
